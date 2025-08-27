from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from scraper.rozee_scraper import scrape_rozee_jobs_selenium
from matcher.resume_matcher import extract_text_from_pdf, match_resume_with_jobs, extract_keywords_text
from models import db, User
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
import pandas as pd
import os
import time
from flask import send_from_directory
from werkzeug.utils import secure_filename
import google.generativeai as genai
from dotenv import load_dotenv
import os

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)


app = Flask(__name__)
app.config["SECRET_KEY"] = "your-secret-key"
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///users.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = "data"

# Initialize extensions
db.init_app(app)
login_manager = LoginManager(app)
login_manager.login_view = "login"

# Global variables
job_cache = []
job_cache_dict = {}  # {job_title: {"timestamp": ..., "df": ...}}
CACHE_EXPIRY_SECONDS = 1800  # 30 min

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ------------------ Auth Routes ------------------

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        email = request.form["email"]
        password = generate_password_hash(request.form["password"])

        user = User(username=username, email=email, password=password)
        db.session.add(user)
        db.session.commit()
        flash("Registration successful. Please log in.", "success")
        return redirect(url_for("login"))

    return render_template("auth/register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]

        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for("dashboard"))
        else:
            flash("Invalid credentials", "danger")

    return render_template("auth/login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))

# ------------------ Helper ------------------

def get_cached_or_scrape_jobs(job_title, pages=2):
    now = time.time()
    cached = job_cache_dict.get(job_title.lower())

    if cached and (now - cached["timestamp"] < CACHE_EXPIRY_SECONDS):
        print(f"✅ Using cached jobs for: {job_title}")
        return cached["df"]

    df = scrape_rozee_jobs_selenium(job_title, pages=pages)
    job_cache_dict[job_title.lower()] = {"timestamp": now, "df": df}
    return df

# ------------------ Main Routes ------------------
@app.route('/')
@login_required
def dashboard():
    return render_template('dashboard.html')

@app.route("/index", methods=["GET", "POST"])
@login_required
def index():
    global job_cache

    if request.method == "POST":
        selected_resume_id = request.form.get("selected_resume")
        resume_file = request.files.get("resume")

        # Upload new resume if user uploaded one
        if resume_file and resume_file.filename.endswith(".pdf"):
            filename = secure_filename(resume_file.filename)
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            resume_file.save(filepath)

            new_resume = Resume(filename=filename, filepath=filepath, user_id=current_user.id)
            db.session.add(new_resume)
            db.session.commit()

            resume_path = filepath

        elif selected_resume_id:
            selected = Resume.query.filter_by(id=selected_resume_id, user_id=current_user.id).first()
            if not selected:
                flash("Invalid resume selected.", "danger")
                return redirect(url_for("index"))
            resume_path = selected.filepath

        else:
            flash("Please upload or select a resume.", "warning")
            return redirect(url_for("index"))

        # Job title & filters
        job_title = request.form["job_title"]
        location_filter = request.form.get("location", "").strip()
        min_score_input = request.form.get("min_score", "")
        min_score = float(min_score_input) if min_score_input.strip() != "" else 0.0

        job_df = get_cached_or_scrape_jobs(job_title, pages=2)
        resume_text = extract_text_from_pdf(resume_path)
        top_matches = match_resume_with_jobs(resume_text, jobs_df=job_df, top_n=100)

        if location_filter:
            top_matches = top_matches[top_matches["location"].str.contains(location_filter, case=False, na=False)]
        top_matches = top_matches[top_matches["match_score"] >= min_score]

        top_matches.to_csv("data/matched_jobs.csv", index=False)
        job_cache = top_matches.to_dict(orient="records")

        return render_template("results.html", matches=job_cache[:10], current_page=1, total_pages=(len(job_cache)+9)//10)

    # GET method — fetch resumes to show in dropdown
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    return render_template("index.html", resumes=resumes)


from models import Resume
from werkzeug.utils import secure_filename
from flask import flash

@app.route("/auto", methods=["GET", "POST"])
@login_required
def auto_mode():
    global job_cache

    resumes = Resume.query.filter_by(user_id=current_user.id).all()

    if request.method == "POST":
        selected_resume_id = request.form.get("selected_resume")
        resume_file = request.files.get("resume")
        resume_path = None

        # Upload new resume if user uploaded one
        if resume_file and resume_file.filename.endswith(".pdf"):
            filename = secure_filename(resume_file.filename)
            resume_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            resume_file.save(resume_path)

            new_resume = Resume(filename=filename, filepath=resume_path, user_id=current_user.id)
            db.session.add(new_resume)
            db.session.commit()

        # Or use existing selected resume
        elif selected_resume_id:
            selected = Resume.query.filter_by(id=selected_resume_id, user_id=current_user.id).first()
            if not selected:
                flash("Invalid resume selected.", "danger")
                return redirect(url_for("auto_mode"))
            resume_path = selected.filepath

        else:
            flash("Please upload or select a resume.", "warning")
            return redirect(url_for("auto_mode"))

        # Resume text + job role detection
        resume_text = extract_text_from_pdf(resume_path)
        top_keyword = extract_keywords_text(resume_text, num_keywords=1)
        query = top_keyword.strip().replace(" ", "+")
        print("🔍 Auto-detected job title:", query)

        job_df = get_cached_or_scrape_jobs(query, pages=2)
        top_matches = match_resume_with_jobs(resume_text, jobs_df=job_df, top_n=100)

        job_cache = top_matches.to_dict(orient="records")
        return redirect(url_for("results_page", page=1))

    return render_template("auto.html", resumes=resumes)


############################################################################################


@app.route("/results")
@login_required
def results_page():
    global job_cache
    page = int(request.args.get("page", 1))
    per_page = 10
    start = (page - 1) * per_page
    end = start + per_page
    total_pages = (len(job_cache) + per_page - 1) // per_page

    return render_template(
        "results.html",
        matches=job_cache[start:end],
        current_page=page,
        total_pages=total_pages
    )

@app.route("/job/<int:job_id>")
@login_required
def job_detail(job_id):
    if 0 <= job_id < len(job_cache):
        job = job_cache[job_id]
        return render_template("job_detail.html", job=job, job_id=job_id)
    return "Job not found", 404

@app.route("/download")
@login_required
def download_csv():
    return send_file("data/matched_jobs.csv", as_attachment=True)


# Resume uploads 
from models import Resume

@app.route("/resumes", methods=["GET", "POST"])
@login_required
def resume_list():
    if request.method == "POST":
        file = request.files["resume"]
        if file and file.filename.endswith(".pdf"):
            filename = f"{current_user.id}_{int(time.time())}_{file.filename}"
            save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(save_path)

            new_resume = Resume(
                user_id=current_user.id,
                filename=filename,
                original_filename=file.filename
            )
            db.session.add(new_resume)
            db.session.commit()
            flash("Resume uploaded successfully!", "success")
            return redirect(url_for("resume_list"))

    user_resumes = Resume.query.filter_by(user_id=current_user.id).all()
    return render_template("resumes.html", resumes=user_resumes)

#delete resume
@app.route("/resume/<int:resume_id>/delete", methods=["POST"])
@login_required
def delete_resume(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    if not resume:
        flash("Resume not found.", "danger")
        return redirect(url_for("index"))

    try:
        if os.path.exists(resume.filepath):
            os.remove(resume.filepath)
        db.session.delete(resume)
        db.session.commit()
        flash("Resume deleted successfully.", "success")
    except Exception as e:
        flash("Error deleting resume.", "danger")

    return redirect(url_for("index"))

#download resume

@app.route("/resume/<int:resume_id>/download")
@login_required
def download_resume(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    if resume:
        return send_from_directory(
            directory=os.path.dirname(resume.filepath),
            path=os.path.basename(resume.filepath),
            as_attachment=True
        )
    flash("Resume not found.", "danger")
    return redirect(url_for("index"))



@app.route("/resume/<int:resume_id>/preview")
@login_required
def preview_resume(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first()
    if resume:
        return send_from_directory(
            directory=os.path.dirname(resume.filepath),
            path=os.path.basename(resume.filepath)
        )
    flash("Resume not found.", "danger")
    return redirect(url_for("index"))

#######################################################################################################

from models import MonitoringRule, WatchlistMatch

@app.route("/monitor", methods=["GET", "POST"])
@login_required
def monitor_rules():
    resumes = Resume.query.filter_by(user_id=current_user.id).all()

    if request.method == "POST":
        job_title = request.form.get("job_title", "").strip()
        resume_id = request.form.get("resume_id", "").strip()

        if not job_title or not resume_id:
            flash("Please enter a job title and select a resume.", "warning")
        else:
            # Prevent duplicate rules
            existing = MonitoringRule.query.filter_by(
                job_title=job_title.lower(),
                resume_id=resume_id,
                user_id=current_user.id
            ).first()

            if existing:
                flash("This monitoring rule already exists.", "info")
            else:
                new_rule = MonitoringRule(
                    job_title=job_title.lower(),
                    resume_id=int(resume_id),
                    user_id=current_user.id
                )
                db.session.add(new_rule)
                db.session.commit()
                flash("✅ Monitoring rule added!", "success")
            return redirect(url_for("monitor_rules"))

    rules = MonitoringRule.query.filter_by(user_id=current_user.id).order_by(MonitoringRule.created_at.desc()).all()
    return render_template("monitor.html", resumes=resumes, rules=rules)


from monitoring.background_job import monitor_jobs_loop
from models import MonitoringRule, WatchlistMatch

monitor_jobs_loop(
    app=app,
    db=db,
    MonitoringRule=MonitoringRule,
    WatchlistMatch=WatchlistMatch,
    scrape_jobs_func=scrape_rozee_jobs_selenium,
    extract_text_func=extract_text_from_pdf,
    match_func=match_resume_with_jobs
)

@app.route("/monitor/delete/<int:rule_id>", methods=["POST"])
@login_required
def delete_monitoring_rule(rule_id):
    rule = MonitoringRule.query.get_or_404(rule_id)

    if rule.user_id != current_user.id:
        flash("Unauthorized action.", "danger")
        return redirect(url_for("monitor_rules"))

    # Optional: delete any watchlist matches related to the rule
    WatchlistMatch.query.filter_by(rule_id=rule.id).delete()

    db.session.delete(rule)
    db.session.commit()
    flash("Monitoring rule deleted successfully.", "success")
    return redirect(url_for("monitor_rules"))


######################watchlist 
@app.route("/watchlist")
@login_required
def watchlist():
    page = request.args.get("page", 1, type=int)
    per_page = 10
    matches = WatchlistMatch.query.join(MonitoringRule).filter(
        MonitoringRule.user_id == current_user.id
    ).order_by(WatchlistMatch.found_at.desc()).paginate(page=page, per_page=per_page)

    return render_template("watchlist.html", matches=matches)
@app.route("/watchlist/delete/<int:match_id>", methods=["POST"])
@login_required
def delete_watchlist_entry(match_id):
    match = WatchlistMatch.query.get_or_404(match_id)

    # Ensure the match belongs to the current user
    if match.monitoring_rule.user_id != current_user.id:
        flash("Unauthorized", "danger")
        return redirect(url_for("watchlist"))

    db.session.delete(match)
    db.session.commit()
    flash("Watchlist entry deleted.", "info")
    return redirect(url_for("watchlist"))

# AI resume feedback ########################3


from models import Resume
from matcher.resume_matcher import extract_text_from_pdf
from flask import request, render_template, redirect, url_for, flash
from flask_login import login_required, current_user
import google.generativeai as genai
import time, os
from flask import session

@app.route("/feedback", methods=["GET", "POST"])
@login_required
def ai_feedback():
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    feedback = None
    improved_resume = None

    if request.method == "POST":
        selected_resume_id = request.form.get("selected_resume")
        uploaded_file = request.files.get("resume")
        resume_path = None
        temp_file = False

        # Upload resume temporarily
        if uploaded_file and uploaded_file.filename.endswith(".pdf"):
            filename = f"temp_{current_user.id}_{int(time.time())}_{uploaded_file.filename}"
            resume_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            uploaded_file.save(resume_path)
            temp_file = True

        elif selected_resume_id:
            selected = Resume.query.filter_by(id=selected_resume_id, user_id=current_user.id).first()
            if not selected:
                flash("Invalid resume selected.", "danger")
                return redirect(url_for("ai_feedback"))
            resume_path = selected.filepath

        else:
            flash("Please upload or select a resume.", "warning")
            return redirect(url_for("ai_feedback"))

        resume_text = extract_text_from_pdf(resume_path)

        # Prompt for feedback
        feedback_prompt = f"""
        You are a professional resume reviewer.
        Please give detailed, constructive feedback on the following resume.
        Focus on formatting, clarity, keywords, structure, and tone.

        Resume:
        {resume_text}
        """

        # Prompt for rewriting
        rewrite_prompt = f"""
        You are an expert resume editor.
        Rewrite and improve the following resume to be ATS-friendly, concise, and impactful.
        Use professional formatting and highlight achievements. Do not add fake experience.

        Resume to improve:
        {resume_text}
        """

        try:
            model = genai.GenerativeModel("models/gemini-1.5-flash")

            # 1️⃣ Generate feedback
            feedback_response = model.generate_content(feedback_prompt)
            feedback = getattr(feedback_response, "text", None)

            # 2️⃣ Generate rewritten resume
            rewrite_response = model.generate_content(rewrite_prompt)
            improved_resume = getattr(rewrite_response, "text", None)

            # Store improved resume in session for download
            session["improved_resume"] = improved_resume

        except Exception as e:
            flash(f"Error generating feedback: {str(e)}", "danger")

        # Delete temp file
        if temp_file and resume_path and os.path.exists(resume_path):
            os.remove(resume_path)

    return render_template("feedback.html", resumes=resumes, feedback=feedback, improved_resume=improved_resume)

from flask import make_response

from docx import Document
from flask import send_file, session
import io

@app.route("/download-improved-resume")
@login_required
def download_improved_resume():
    if "improved_resume" not in session:
        flash("No improved resume found to download.", "warning")
        return redirect(url_for("ai_feedback"))

    content = session["improved_resume"]

    # Create a Word document in memory
    doc = Document()
    doc.add_heading("Improved Resume", 0)

    for line in content.strip().split("\n"):
        doc.add_paragraph(line.strip())

    # Save document to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="Improved_Resume.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# desciption matching
@app.route("/description-match", methods=["GET", "POST"])
@login_required
def description_match():
    feedback = None
    score = None

    if request.method == "POST":
        job_description = request.form.get("job_description")
        uploaded_file = request.files.get("resume")

        if not job_description or not uploaded_file:
            flash("Please provide both job description and a resume.", "warning")
            return redirect(url_for("description_match"))

        # Save the uploaded resume
        filename = f"{current_user.id}_{int(time.time())}_{uploaded_file.filename}"
        resume_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        uploaded_file.save(resume_path)

        # Extract text from resume
        resume_text = extract_text_from_pdf(resume_path)

        # Create prompt
        prompt = f"""
Compare the following resume content to the job description.

Give a match score between 0 and 100 based on alignment of skills, experience, and keywords.

Then, give a short suggestion: should this candidate apply or not?

Resume:
{resume_text}

Job Description:
{job_description}
"""

        try:
            model = genai.GenerativeModel("gemini-1.5-flash")  # or the latest working free model
            response = model.generate_content(prompt)
            if hasattr(response, "text"):
                feedback = response.text
                # Try extracting a score if model includes it in response
                import re
                match = re.search(r"(\d{1,3})\s*\/?\s*100", feedback)
                if match:
                    score = int(match.group(1))
        except Exception as e:
            flash(f"Error generating match score: {str(e)}", "danger")

    return render_template("description_match.html", feedback=feedback, score=score)



# from flask import jsonify
# import re
# from matcher.resume_matcher import extract_text_from_pdf

# @app.route("/resume-feedback", methods=["GET", "POST"])
# @login_required
# def resume_feedback():
#     feedback = ""
#     resume_text = ""

#     if request.method == "POST":
#         resume_file = request.files.get("resume")
#         if resume_file and resume_file.filename.endswith(".pdf"):
#             filename = secure_filename(resume_file.filename)
#             path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
#             resume_file.save(path)

#             resume_text = extract_text_from_pdf(path)
#             feedback = analyze_resume_text(resume_text)

#     return render_template("resume_feedback.html", feedback=feedback, text=resume_text)

# def analyze_resume_text(text):
#     suggestions = []

#     if len(text) < 300:
#         suggestions.append("📝 Your resume seems too short. Consider adding more details about your experience.")

#     if "objective" not in text.lower() and "summary" not in text.lower():
#         suggestions.append("📌 Missing a professional summary or objective section.")

#     if "experience" not in text.lower():
#         suggestions.append("🔍 No 'Experience' section found. Add work experience to strengthen your resume.")

#     if "skills" not in text.lower():
#         suggestions.append("🧠 No 'Skills' section detected. List your technical or soft skills clearly.")

#     if re.search(r'\b(c++, java, python, html, css, sql, react)\b', text.lower()) is None:
#         suggestions.append("💡 Consider adding relevant keywords like programming languages or tools.")

#     if not suggestions:
#         return "✅ Your resume looks good! All key sections are present."
    
#     return "\n".join(suggestions)



# ------------------ Run App ------------------

if __name__ == "__main__":
    app.run(debug=True)
